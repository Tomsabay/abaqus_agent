"""Assemble a Word calculation report (计算书) from one finished run directory.

BUILD_SPEC W20 criterion 1. Usage:

    .venv\\Scripts\\python.exe reporting\\build_docx_report.py <run_dir> \\
        --template course\\materials\\module5\\report_template.docx \\
        --limits   course\\materials\\module5\\limits.yaml \\
        --out      artifacts\\course\\5_3\\计算书.docx

Design rules that are not negotiable:

* the verdict comes from the limit file, computed by ``reporting.limits`` — with
  no ``--limits`` the conclusion section is *refused*, never guessed;
* the cover's Abaqus release comes from ``abaqus information=release``, never from
  ``spec.meta.abaqus_release`` (unvalidated metadata that has drifted before);
* figures are embedded byte-for-byte, so a picture inside the .docx hashes the
  same as the PNG the solver produced;
* a run with no solver behind it prints the no-solve banner on the first screen
  and carries no numeric KPI table.

Exit codes: 0 ok · 2 bad input/limits · 4 template has placeholders this tool
does not know how to fill.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import sys
import zipfile
from pathlib import Path
from typing import Any

if __package__ in (None, ""):  # direct `python reporting/build_docx_report.py`
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402
from docx.shared import Inches  # noqa: E402

from reporting.limits import (  # noqa: E402
    NO_LIMITS_NOTICE,
    Judgement,
    LimitsError,
    conclusion_lines,
    format_number,
    format_pct,
    judge_kpis,
    load_limits,
    max_utilisation,
    overall_verdict,
)
from reporting.run_bundle import (  # noqa: E402
    RunBundle,
    actual_abaqus_release,
    generated_at,
    load_run_bundle,
    no_solve_notice,
)
from reporting.templates import NO_SOLVE_BANNER  # noqa: E402

FIGURE_WIDTH_INCHES = 5.8
MAX_FIGURES = 3
MAX_ARTIFACT_ROWS = 14
NO_FIGURE_TEXT = "（本次运行未产出云图/曲线图，图位留空）"
NO_SOLVE_KPI_ROW = f"（{NO_SOLVE_BANNER} —— 未生成任何数值 KPI）"


# ── docx helpers ──────────────────────────────────────────────────────

def _all_paragraphs(document) -> list:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    return paragraphs


def _replace_inline(document, mapping: dict[str, str]) -> set[str]:
    """Replace inline placeholders, tolerating runs that split a placeholder."""
    used: set[str] = set()
    for paragraph in _all_paragraphs(document):
        text = paragraph.text
        if "{{" not in text:
            continue
        new_text = text
        for placeholder, value in mapping.items():
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, value)
                used.add(placeholder)
        if new_text == text:
            continue
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)
    return used


def _find_block_paragraph(document, placeholder: str):
    for paragraph in _all_paragraphs(document):
        if paragraph.text.strip() == placeholder:
            return paragraph
    return None


def _drop(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _insert_lines_before(paragraph, lines: list[str]) -> None:
    for line in lines:
        paragraph.insert_paragraph_before(line)


def _insert_table_before(document, paragraph, header: list[str],
                         rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, header):
        cell.text = title
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
    paragraph._p.addprevious(table._tbl)


def _insert_pictures_before(paragraph, images: list[Path]) -> list[Path]:
    embedded: list[Path] = []
    for image in images[:MAX_FIGURES]:
        holder = paragraph.insert_paragraph_before()
        try:
            holder.add_run().add_picture(str(image), width=Inches(FIGURE_WIDTH_INCHES))
        except Exception as e:  # unreadable/unsupported image must not kill the report
            _drop(holder)
            paragraph.insert_paragraph_before(f"（图片 {image.name} 无法嵌入：{e}）")
            continue
        paragraph.insert_paragraph_before(f"图 {len(embedded) + 1}  {image.name}")
        embedded.append(image)
    if not embedded:
        paragraph.insert_paragraph_before(NO_FIGURE_TEXT)
    return embedded


# ── content builders ──────────────────────────────────────────────────

def _kpi_descriptions(bundle: RunBundle) -> dict[str, str]:
    described: dict[str, str] = {}
    for item in bundle.result.get("odb_lens_recipe", []) or []:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", ""))
        if not name:
            continue
        parts = [str(item.get("type", ""))]
        for key in ("variable", "field_variable", "component", "location"):
            if item.get(key):
                parts.append(f"{key}={item[key]}")
        described[name] = " / ".join(p for p in parts if p)
    return described


def _v2_parts_summary(spec: dict) -> str:
    """One phrase per part: what it is dimensionally, and how finely it is meshed.

    v1 had a single `geometry` block whose scalars were the whole story. v2 has
    N parts each with its own mesh, so the seed has to be attributed to a part
    or it means nothing.
    """
    described = []
    for part in spec.get("parts") or []:
        if not isinstance(part, dict):
            continue
        bits = [str(part.get("name") or "?")]
        if part.get("dimensionality"):
            bits.append(str(part["dimensionality"]))
        mesh = part.get("mesh") if isinstance(part.get("mesh"), dict) else {}
        if mesh.get("seed") is not None:
            bits.append(f"seed={mesh['seed']}")
        if mesh.get("element"):
            bits.append(str(mesh["element"]))
        described.append(" ".join(bits))
    return "；".join(described) if described else "（未记录）"


def _v2_conditions_summary(spec: dict) -> str:
    """The Abaqus call each condition makes, with the step it is created in.

    Deliberately the call name and not a prettified label: `Pressure` vs
    `SurfaceTraction` is the difference between a load along the surface normal
    and one along a stated vector, and a report that blurred them would be
    describing a different analysis than the one that ran.
    """
    described = []
    for cond in spec.get("conditions") or []:
        if not isinstance(cond, dict):
            continue
        call = str(cond.get("call") or "?")
        magnitude = cond.get("magnitude")
        if magnitude is not None:
            call += f"({magnitude})"
        described.append(call)
    return "、".join(described) if described else "（未记录）"


def _spec_summary_lines(bundle: RunBundle, release: str) -> list[str]:
    spec = bundle.spec if isinstance(bundle.spec, dict) else {}
    meta = spec.get("meta", {}) if isinstance(spec.get("meta"), dict) else {}
    material = spec.get("material", {}) if isinstance(spec.get("material"), dict) else {}
    analysis = spec.get("analysis", {}) if isinstance(spec.get("analysis"), dict) else {}
    bc_load = spec.get("bc_load", {}) if isinstance(spec.get("bc_load"), dict) else {}
    geometry = spec.get("geometry", {}) if isinstance(spec.get("geometry"), dict) else {}

    lines = []
    if meta.get("description"):
        lines.append(f"工程/算例描述：{meta['description']}")
    # `geometry` / `bc_load` are v1's, removed 2026-08-16. The reads stay
    # because reports are built from ARCHIVED run directories, and the v1 spec
    # frozen inside one of those must still render. New runs take the v2 branch
    # below; without it the report silently dropped both lines, which reads as
    # "the model had no geometry and no loads" rather than as a missing feature.
    if geometry:
        geo = "、".join(f"{k}={v}" for k, v in geometry.items() if not isinstance(v, (dict, list)))
        if geo:
            lines.append(f"几何与网格参数：{geo}")
    elif spec.get("parts"):
        lines.append(f"几何与网格参数：{_v2_parts_summary(spec)}")
    elif isinstance(spec.get("deck"), dict):
        lines.append(
            f"几何与网格参数：整份 .inp 原样提交（{spec['deck'].get('file') or '?'}），"
            "零件/网格/边界条件均由该文件定义")
    if material:
        mat = "、".join(f"{k}={v}" for k, v in material.items() if not isinstance(v, (dict, list)))
        lines.append(f"材料参数：{mat}")
    if analysis:
        ana = "、".join(f"{k}={v}" for k, v in analysis.items() if not isinstance(v, (dict, list)))
        lines.append(f"分析设置：{ana}")
    if bc_load:
        load = "、".join(f"{k}={v}" for k, v in bc_load.items() if not isinstance(v, (dict, list)))
        lines.append(f"边界与荷载：{load}")
    elif spec.get("conditions"):
        lines.append(f"边界与荷载：{_v2_conditions_summary(spec)}")
    if meta.get("units"):
        lines.append(f"单位制：{meta['units']}")
    lines.append(f"Problem Spec 文件：{bundle.spec_path or '（未记录）'}")
    lines.append(f"求解器版本（实测 abaqus information=release）：{release}")
    spec_release = bundle.spec_release
    if spec_release and spec_release != release:
        lines.append(
            f"版本口径核对：spec 声明 {spec_release}，与实测 {release} 不一致；"
            "本报告一律以实测版本为准（spec 中该字段不参与命令拼装）。"
        )
    if not lines:
        lines.append("（未找到 Problem Spec，工程概况留空）")
    return lines


def _model_table(bundle: RunBundle) -> tuple[list[str], list[list[str]]]:
    rows: list[list[str]] = []
    inp = bundle.result.get("stages", {}).get("build_model", {}) if bundle.result else {}
    if isinstance(inp, dict) and inp.get("inp_path"):
        rows.append(["INP 文件", str(inp["inp_path"])])
    if bundle.mesh_counts.get("elements"):
        rows.append(["单元数", f"{bundle.mesh_counts['elements']}（来自 {Path(bundle.dat_path).name}）"])
    if bundle.mesh_counts.get("nodes"):
        rows.append(["节点数", f"{bundle.mesh_counts['nodes']}（来自 {Path(bundle.dat_path).name}）"])
    submit = bundle.result.get("stages", {}).get("submit_job", {}) if bundle.result else {}
    if isinstance(submit, dict):
        if submit.get("cpus"):
            rows.append(["求解 CPU 数", str(submit["cpus"])])
        if submit.get("mp_mode"):
            rows.append(["并行模式", str(submit["mp_mode"])])
    analysis = bundle.spec.get("analysis", {}) if isinstance(bundle.spec, dict) else {}
    if isinstance(analysis, dict):
        if analysis.get("solver"):
            rows.append(["求解器", str(analysis["solver"])])
        if analysis.get("step_type"):
            rows.append(["分析步类型", str(analysis["step_type"])])
    # v2 names the step by the Abaqus method it calls, and can declare several.
    # Reading only `analysis.step_type` left this row blank on every new run.
    if isinstance(bundle.spec, dict) and not analysis.get("step_type"):
        calls = [str(s.get("call")) for s in (bundle.spec.get("steps") or [])
                 if isinstance(s, dict) and s.get("call")]
        if calls:
            rows.append(["分析步类型", " → ".join(calls)])
    if bundle.started_at:
        rows.append(["求解开始", bundle.started_at])
    if bundle.finished_at:
        rows.append(["求解结束", bundle.finished_at])
    if not rows:
        rows.append(["（无模型信息）", "-"])
    return ["项", "值"], rows


def _kpi_table(bundle: RunBundle, descriptions: dict[str, str],
               judgements: list[Judgement]) -> tuple[list[str], list[list[str]]]:
    header = ["序号", "KPI", "提取口径", "数值", "单位", "数据来源文件"]
    if bundle.unsolved or not bundle.kpis:
        return header, [["-", NO_SOLVE_KPI_ROW, "-", "-", "-", "-"]]
    unit_by_name = {j.name: j.unit for j in judgements}
    source_name = Path(bundle.kpi_source).name if bundle.kpi_source else "-"
    rows = []
    for index, (name, value) in enumerate(bundle.kpis.items(), start=1):
        rows.append([
            str(index),
            name,
            descriptions.get(name, "-"),
            format_number(value),
            unit_by_name.get(name) or "-",
            source_name,
        ])
    return header, rows


def _limit_table(judgements: list[Judgement]) -> tuple[list[str], list[list[str]]]:
    header = ["KPI", "说明", "实测值", "限值", "单位", "关系", "利用率", "判定", "限值来源"]
    rows = []
    for j in judgements:
        if j.limit is None:
            rows.append([j.name, "-", format_number(j.actual), "-", "-", "-", "-", j.verdict, "-"])
            continue
        rows.append([
            j.name,
            j.label,
            format_number(j.compared),
            format_number(j.limit.value),
            j.unit,
            j.limit.relation,
            format_pct(j.utilisation_pct),
            j.verdict,
            j.source,
        ])
    if not rows:
        rows.append(["-", "-", "-", "-", "-", "-", "-", "-", "-"])
    return header, rows


def _artifact_table(bundle: RunBundle) -> tuple[list[str], list[list[str]]]:
    header = ["产物", "字节数", "sha256（前 16 位）"]
    rows = []
    for name in sorted(bundle.artifacts):
        entry = bundle.artifacts[name]
        if not isinstance(entry, dict):
            continue
        rows.append([
            name,
            str(entry.get("bytes", "-")),
            str(entry.get("sha256", ""))[:16] or "-",
        ])
    if len(rows) > MAX_ARTIFACT_ROWS:
        remaining = len(rows) - MAX_ARTIFACT_ROWS
        rows = rows[:MAX_ARTIFACT_ROWS]
        rows.append([f"（另有 {remaining} 项产物，完整清单见 capsule.json）", "-", "-"])
    if not rows:
        rows.append(["（capsule.json 未记录产物）", "-", "-"])
    return header, rows


def _repro_command(bundle: RunBundle) -> str:
    if not bundle.spec_path:
        return "（未记录 spec 路径，无法给出复现命令）"
    spec = Path(bundle.spec_path)
    case_dir = spec.parent
    expected = case_dir / "expected.json"
    runner = case_dir / "runner.json"
    parts = [".venv\\Scripts\\python.exe", "agent\\orchestrator.py", str(spec)]
    if expected.is_file():
        parts.append(str(expected))
    if runner.is_file():
        parts.append(str(runner))
    return " ".join(parts)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1 << 20), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _docx_media_hashes(path: Path) -> dict[str, str]:
    hashes = {}
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/"):
                hashes[name] = hashlib.sha256(archive.read(name)).hexdigest()
    return hashes


# ── main assembly ─────────────────────────────────────────────────────

def build_report(
    run_dir: str | Path,
    template: str | Path,
    out: str | Path,
    *,
    limits_path: str | Path | None = None,
    spec_path: str | Path | None = None,
    title: str | None = None,
    report_no: str | None = None,
    release_timeout: float = 90.0,
) -> dict[str, Any]:
    """Fill a template with one run's data and write the .docx. Returns a summary."""
    bundle = load_run_bundle(run_dir, spec_path=spec_path)
    template_path = Path(template)
    if not template_path.is_file():
        raise FileNotFoundError(f"模板不存在：{template_path}")

    limits = load_limits(limits_path) if limits_path else {}
    judgements = judge_kpis(bundle.kpis, limits) if bundle.kpis else []
    release = actual_abaqus_release(timeout=release_timeout)
    descriptions = _kpi_descriptions(bundle)

    document = Document(str(template_path))

    inline = {
        "{{PROJECT_TITLE}}": title or f"{bundle.model_name} 有限元计算书",
        "{{MODEL_NAME}}": bundle.model_name,
        "{{REPORT_NO}}": report_no or f"CALC-{bundle.run_id}",
        "{{RUN_ID}}": bundle.run_id,
        "{{ABAQUS_RELEASE}}": release,
        "{{SOLVE_STATUS}}": bundle.status,
        "{{GENERATED_AT}}": generated_at(),
        "{{ODB_PATH}}": bundle.odb_path or "（未记录 ODB 路径）",
        "{{REPRO_COMMAND}}": _repro_command(bundle),
    }
    _replace_inline(document, inline)

    embedded: list[Path] = []
    filled_blocks: list[str] = []

    banner_paragraph = _find_block_paragraph(document, "{{NO_SOLVE_BANNER}}")
    if banner_paragraph is not None:
        if bundle.unsolved:
            banner_paragraph.text = f"{NO_SOLVE_BANNER} —— {no_solve_notice(bundle)}"
            for run in banner_paragraph.runs:
                run.bold = True
        else:
            _drop(banner_paragraph)
        filled_blocks.append("{{NO_SOLVE_BANNER}}")

    block_specs: list[tuple[str, str, Any]] = [
        ("{{SPEC_SUMMARY}}", "lines", _spec_summary_lines(bundle, release)),
        ("{{MODEL_TABLE}}", "table", _model_table(bundle)),
        ("{{KPI_TABLE}}", "table", _kpi_table(bundle, descriptions, judgements)),
        ("{{FIGURE}}", "pictures", bundle.images),
        (
            "{{LIMIT_TABLE}}",
            "table" if limits_path else "lines",
            _limit_table(judgements) if limits_path else [NO_LIMITS_NOTICE],
        ),
        (
            "{{CONCLUSION}}",
            "lines",
            conclusion_lines(judgements, limits_path if limits_path else None),
        ),
        ("{{ARTIFACT_TABLE}}", "table", _artifact_table(bundle)),
    ]
    for placeholder, kind, payload in block_specs:
        paragraph = _find_block_paragraph(document, placeholder)
        if paragraph is None:
            continue
        if kind == "lines":
            _insert_lines_before(paragraph, list(payload))
        elif kind == "table":
            header, rows = payload
            _insert_table_before(document, paragraph, header, rows)
        elif kind == "pictures":
            embedded = _insert_pictures_before(paragraph, list(payload))
        _drop(paragraph)
        filled_blocks.append(placeholder)

    leftovers = sorted({
        "{{" + chunk.split("}}")[0] + "}}"
        for paragraph in _all_paragraphs(document)
        for chunk in paragraph.text.split("{{")[1:]
        if "}}" in chunk
    })
    if leftovers:
        raise UnknownPlaceholderError(leftovers)

    out_path = Path(out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))

    media = _docx_media_hashes(out_path)
    worst = max_utilisation(judgements)
    return {
        "output_path": str(out_path),
        "bytes": out_path.stat().st_size,
        "template": str(template_path),
        "run_dir": str(bundle.run_dir),
        "run_id": bundle.run_id,
        "model_name": bundle.model_name,
        "status": bundle.status,
        "unsolved": bundle.unsolved,
        "abaqus_release": release,
        "spec_path": bundle.spec_path,
        "spec_declared_release": bundle.spec_release,
        "kpi_source": bundle.kpi_source,
        "kpis": bundle.kpis,
        "kpi_values_formatted": {k: format_number(v) for k, v in bundle.kpis.items()},
        "limits_path": str(limits_path) if limits_path else "",
        "conclusion_section": bool(limits_path),
        "overall_verdict": overall_verdict(judgements) if limits_path else "",
        "max_utilisation_pct": worst.utilisation_pct if worst else None,
        "judgements": [
            {
                "name": j.name,
                "actual": j.actual,
                "compared": j.compared,
                "limit": j.limit.value if j.limit else None,
                "unit": j.unit,
                "utilisation_pct": j.utilisation_pct,
                "verdict": j.verdict,
            }
            for j in judgements
        ],
        "embedded_images": [
            {"name": image.name, "sha256": _sha256(image), "bytes": image.stat().st_size}
            for image in embedded
        ],
        "docx_media_sha256": media,
        "filled_blocks": filled_blocks,
    }


class UnknownPlaceholderError(ValueError):
    """The template contains {{...}} markers this tool cannot fill."""

    def __init__(self, placeholders: list[str]) -> None:
        self.placeholders = placeholders
        super().__init__(
            "模板含本工具无法识别的占位符（未填充，已中止）："
            + "、".join(placeholders)
        )


def docx_text(path: str | Path) -> str:
    """All text of a generated report (body + tables + header/footer)."""
    document = Document(str(path))
    chunks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in document.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(chunks)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="build_docx_report.py",
        description="Assemble a Word calculation report from a finished run directory",
    )
    parser.add_argument("run_dir", help="Run directory (contains result.json / _kpi_result.json)")
    parser.add_argument("--template", required=True, help="report_template.docx path")
    parser.add_argument("--out", required=True, help="Output .docx path")
    parser.add_argument(
        "--limits",
        help="Limit table YAML. Omit it and the conclusion section is refused, not guessed.",
    )
    parser.add_argument("--spec", help="Override the Problem Spec path recorded in result.json")
    parser.add_argument("--title", help="Cover title (default: <model> 有限元计算书)")
    parser.add_argument("--report-no", help="Report number on the cover")
    parser.add_argument("--release-timeout", type=float, default=90.0)
    parser.add_argument("--json", action="store_true", dest="as_json", help="Print JSON summary")
    args = parser.parse_args(argv)

    try:
        summary = build_report(
            args.run_dir,
            args.template,
            args.out,
            limits_path=args.limits,
            spec_path=args.spec,
            title=args.title,
            report_no=args.report_no,
            release_timeout=args.release_timeout,
        )
    except (FileNotFoundError, LimitsError) as e:
        print(f"错误：{e}", file=sys.stderr)
        return 2
    except UnknownPlaceholderError as e:
        print(f"错误：{e}", file=sys.stderr)
        return 4

    if args.as_json:
        print(json.dumps(summary, ensure_ascii=False, indent=2))
    else:
        print(f"已生成计算书：{summary['output_path']}（{summary['bytes']} 字节）")
        print(f"run_id={summary['run_id']}  Abaqus（实测）={summary['abaqus_release']}")
        print(f"KPI 来源：{summary['kpi_source'] or '-'}")
        if summary["conclusion_section"]:
            print(f"限值文件：{summary['limits_path']}  总体判定：{summary['overall_verdict']}"
                  f"  最大利用率：{format_pct(summary['max_utilisation_pct'])}")
        else:
            print(NO_LIMITS_NOTICE)
        for image in summary["embedded_images"]:
            print(f"内嵌图片：{image['name']}  sha256={image['sha256']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
